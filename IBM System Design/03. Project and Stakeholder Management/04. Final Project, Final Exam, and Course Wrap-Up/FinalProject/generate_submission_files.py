import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

base_dir = r"d:\Ent\Learning\Cousera\IBM System Design\03. Project and Stakeholder Management\04. Final Project, Final Exam, and Course Wrap-Up"

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_table(table, col_widths, headers, data):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F497D") # Navy blue
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
    
    # Data rows
    for row_idx, row_data in enumerate(data):
        row_cells = table.add_row().cells
        bg_color = "F2F5F9" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = str(text)
            set_cell_background(row_cells[col_idx], bg_color)
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(9.5)

    # Set column widths
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = Inches(w)

# ==========================================
# Task 1: Business Case Review
# ==========================================
doc1 = Document()
doc1.add_heading("Task 1: Business Case Review - Shop Ease Online AI Chatbot", level=1)
p = doc1.add_paragraph("Below is the business case analysis outlining key Goals, Challenges, Assumptions, Data Gaps, and Clarification Questions for the Shop Ease Online AI Chatbot implementation.")

headers1 = ["Goal", "Challenge", "Assumption", "Data Gap", "Clarification Question"]
data1 = [
    [
        "Reduce customer service operating costs by 30% through automated inquiry resolution.",
        "Surge in customer inquiry volume overwhelming live agents and causing long email response backlogs.",
        "Shop Ease Online's CRM and Order Management Systems have modern, robust REST/GraphQL APIs ready for chatbot integration.",
        "Current volume and breakdown of routine vs. complex inquiries (e.g., % of order tracking vs. technical complaints).",
        "What specific APIs or database access protocols are currently exposed by the CRM and OMS?"
    ],
    [
        "Improve Customer Satisfaction (CSAT) score from 78% to 90%.",
        "Inconsistent service quality and prolonged wait times currently damaging customer trust.",
        "A response time of under 10 seconds and 24/7 availability will directly boost customer satisfaction and loyalty.",
        "Detailed baseline metrics on peak traffic hours, inquiry categories, and root causes of customer dissatisfaction.",
        "What criteria will determine whether an inquiry is resolved successfully by AI vs. transferred to a human agent?"
    ],
    [
        "Achieve response times under 10 seconds with 24/7 automated support availability.",
        "High operational costs associated with staffing live agents 24/7 and during non-business hours.",
        "Customers are receptive to interacting with an AI chatbot for routine transactions like order tracking.",
        "Historical data on multi-language distribution across international and non-English customer segments.",
        "Which primary languages must be supported in Phase 1 vs. subsequent rollout phases?"
    ],
    [
        "Increase sales conversion rates during off-hours by answering pre-purchase questions instantly.",
        "Risk of user frustration and drop-off if complex issues are not handed off seamlessly to human agents.",
        "Live agents can be effectively retrained within 4 weeks to handle complex, escalated customer cases.",
        "Cost baseline per support ticket (live agent vs. email) to accurately measure the 30% cost reduction target.",
        "What is the agreed fallback mechanism and SLA for human handoff during off-hours when live agents are unavailable?"
    ]
]
table1 = doc1.add_table(rows=1, cols=5)
format_table(table1, [1.3, 1.4, 1.4, 1.4, 1.5], headers1, data1)
doc1.save(os.path.join(base_dir, "Task1_BusinessCaseReview.docx"))

# ==========================================
# Task 2: Requirements (MoSCoW)
# ==========================================
doc2 = Document()
doc2.add_heading("Task 2: Requirements Gathering & Prioritization - Shop Ease AI Chatbot", level=1)
doc2.add_paragraph("This document outlines 12 prioritized functional and non-functional requirements utilizing the MoSCoW framework for the Shop Ease Online AI Chatbot solution.")

headers2 = ["Requirement Name", "Type", "Priority (MoSCoW)", "Rationale"]
data2 = [
    # 8 Functional Requirements
    ["24/7 Real-Time Automated Q&A", "Functional", "Must Have", "Provides instant answers to routine inquiries and product questions at any time, reducing baseline ticket volume."],
    ["Real-Time Order Tracking & Status Lookup", "Functional", "Must Have", "Integrates with OMS to allow customers to check real-time order and shipping status without human assistance."],
    ["CRM Integration for Customer History", "Functional", "Must Have", "Enables personalized support by retrieving customer profile, purchase history, and open tickets directly from CRM."],
    ["Seamless Live Agent Escalation & Handoff", "Functional", "Must Have", "Transfers complex or sensitive inquiries to human agents with full conversation context to avoid customer frustration."],
    ["Multi-Language Support (Top 3 Languages)", "Functional", "Should Have", "Broadens accessibility and improves CSAT for non-native speaking and international customer segments."],
    ["Off-Hours Pre-Purchase Recommendation", "Functional", "Should Have", "Suggests relevant electronics products based on customer queries to directly drive off-hours sales conversion."],
    ["Automated Post-Chat CSAT Survey", "Functional", "Should Have", "Collects immediate customer feedback post-interaction to measure and track the target 90% CSAT goal."],
    ["Automated Return & Refund Guidance", "Functional", "Nice to Have", "Guides users step-by-step through self-service return and refund requests, further lowering agent workloads."],
    # 4 Non-Functional Requirements
    ["Sub-10 Second Response Latency", "Non-Functional", "Must Have", "Ensures high responsiveness and smooth conversational experience, directly meeting the core business KPI."],
    ["GDPR & Data Privacy Compliance", "Non-Functional", "Must Have", "Protects sensitive customer PII and payment data, ensuring regulatory adherence and preventing legal liabilities."],
    ["High System Availability (99.9% Uptime)", "Non-Functional", "Should Have", "Guarantees 24/7 reliability and minimal downtime during peak shopping hours and promotional campaigns."],
    ["High Concurrent Scalability (5,000 Users)", "Non-Functional", "Should Have", "Maintains performance stability and fast response times during major traffic spikes such as Black Friday or holiday sales."]
]
table2 = doc2.add_table(rows=1, cols=4)
format_table(table2, [1.8, 1.1, 1.2, 2.9], headers2, data2)
doc2.save(os.path.join(base_dir, "Task2_Requirements.docx"))

# ==========================================
# Task 3: RACI Chart
# ==========================================
doc3 = Document()
doc3.add_heading("Task 3: RACI Matrix - Shop Ease Online AI Chatbot", level=1)
doc3.add_paragraph("This RACI framework clarifies the roles, responsibilities, and engagement levels for all seven key stakeholder groups across the project lifecycle.")

headers3 = ["Stakeholder", "Role Description", "RACI Designator(s)"]
data3 = [
    [
        "Executive Sponsor (CEO)",
        "Provides overarching strategic vision, champions organizational alignment, and holds ultimate executive authority for budget approval and project sign-off.",
        "Accountable (A)"
    ],
    [
        "Program Manager",
        "Leads project planning, coordinates cross-functional workstreams, manages schedule, risks, scope, budget, and overall day-to-day execution.",
        "Responsible (R), Accountable (A)"
    ],
    [
        "IT/AI Lead",
        "Directs technical architecture, AI model configuration, CRM/OMS system integrations, infrastructure performance, and technical testing.",
        "Responsible (R)"
    ],
    [
        "Customer Service Manager",
        "Defines support workflows, knowledge base content, agent handoff protocols, and leads live-agent training and operational adoption.",
        "Responsible (R), Consulted (C)"
    ],
    [
        "Marketing Manager",
        "Develops promotional campaigns, drives customer awareness of the new AI chatbot, and aligns conversational tone with brand voice.",
        "Consulted (C), Informed (I)"
    ],
    [
        "Data Protection Officer (DPO)",
        "Conducts privacy impact assessments, audits data handling practices, and ensures strict compliance with GDPR and relevant privacy regulations.",
        "Consulted (C)"
    ],
    [
        "End Customers",
        "Direct users of the chatbot interface who interact with the system, submit inquiries, and provide CSAT ratings and feedback.",
        "Informed (I)"
    ]
]
table3 = doc3.add_table(rows=1, cols=3)
format_table(table3, [1.8, 4.0, 1.2], headers3, data3)
doc3.save(os.path.join(base_dir, "Task3_RACI.docx"))

# ==========================================
# Task 4: Stakeholder Engagement Strategies
# ==========================================
doc4 = Document()
doc4.add_heading("Task 4: Stakeholder Engagement Strategies - Shop Ease AI Chatbot", level=1)
doc4.add_paragraph("This matrix maps tailored engagement and communication strategies based on the Power/Interest Grid for all key stakeholders.")

headers4 = ["Stakeholder", "Power", "Interest", "Engagement Strategy", "Communication Method", "Communication Frequency"]
data4 = [
    [
        "Executive Sponsor (CEO)",
        "High",
        "High",
        "Manage Closely: Provide strategic progress summaries, key milestone achievements, ROI tracking, and seek rapid decisions on major scope/budget changes.",
        "Executive Steering Meetings & Monthly Executive Briefings / Dashboards",
        "Bi-weekly / Monthly"
    ],
    [
        "Program Manager",
        "High",
        "High",
        "Manage Closely: Drive daily alignment, monitor workstream dependencies, resolve blockers proactively, and track deliverables against schedule.",
        "Daily Standups, Weekly PMO Reviews, Jira/Project Management Dashboards",
        "Daily / Continuous"
    ],
    [
        "IT/AI Lead",
        "Medium",
        "High",
        "Keep Satisfied / Active Partner: Collaborate closely on technical architecture, system integration milestones, API readiness, and security standards.",
        "Sprint Planning, Architecture Review Sessions, Slack/Teams Dedicated Channel",
        "Weekly / As-needed"
    ],
    [
        "Customer Service Manager",
        "Medium",
        "High",
        "Keep Satisfied / Active Partner: Engage continuously on chatbot knowledge base training, escalation workflows, change management, and agent feedback.",
        "Bi-weekly Working Sessions, User Acceptance Testing (UAT) Workshops",
        "Bi-weekly"
    ],
    [
        "Marketing Manager",
        "Low",
        "High",
        "Keep Informed: Keep updated on launch timelines, brand voice alignment, promotional release dates, and key feature rollouts.",
        "Email Newsletters, Marketing Alignment Meetings",
        "Bi-weekly / Milestone-based"
    ],
    [
        "Data Protection Officer (DPO)",
        "High",
        "Medium",
        "Keep Satisfied: Involve proactively in data flow reviews, security audits, privacy policy sign-offs, and compliance checkpoints.",
        "Formal Compliance Review Meetings, Audit Report Sign-offs",
        "Monthly / Milestone Gates"
    ],
    [
        "End Customers",
        "Low",
        "High",
        "Keep Informed: Announce new support capabilities, highlight 24/7 self-service convenience, and capture feedback via post-chat surveys.",
        "In-app Announcements, Website Banners, Email Updates, Chatbot Welcome Greetings",
        "At Launch & Continuous"
    ]
]
table4 = doc4.add_table(rows=1, cols=6)
format_table(table4, [1.3, 0.7, 0.7, 1.8, 1.4, 1.1], headers4, data4)
doc4.save(os.path.join(base_dir, "Task4_StakeholderEngagement.docx"))

# ==========================================
# Task 5: Reflection
# ==========================================
doc5 = Document()
doc5.add_heading("Task 5: Reflection on Generative AI in Program & Stakeholder Management", level=1)

reflection_text = (
    "Generative AI significantly enhanced the efficiency and analytical depth throughout this project. By processing the Shop Ease Online business case, GenAI rapidly generated structured baselines for requirements elicitation, RACI assignment, and stakeholder engagement matrices. It helped overcome the 'blank page syndrome,' allowing me to transition quickly from initial ideation to structured analysis. Furthermore, GenAI provided creative perspectives on potential edge cases, such as off-hours sales conversion and multi-language support, which enriched the functional scope.\n\n"
    "However, several notable limitations were observed in the AI's raw outputs. AI responses often defaulted to generic e-commerce assumptions rather than strictly adhering to the technical and operational boundaries of the scenario. For instance, initial suggestions included complex AI capabilities that exceeded the immediate project timeline or overlooked specific legacy CRM integration constraints. Additionally, AI lacked organizational context regarding internal company politics and actual resource availability, occasionally misallocating stakeholder accountabilities.\n\n"
    "To effectively mitigate these limitations, a Business Analyst or Project Manager must exercise professional judgment through a rigorous validation workflow. First, AI outputs must be cross-referenced against business objectives, contractual constraints, and technical architecture specifications. Second, requirements and RACI matrices should be validated through collaborative stakeholder workshops (e.g., UAT sessions and compliance reviews with the DPO and IT Lead). Finally, iterative prompt engineering and human-in-the-loop review are essential to refine AI-generated drafts into actionable, production-ready deliverables that ensure true business alignment and risk governance."
)

for paragraph in reflection_text.split("\n\n"):
    p = doc5.add_paragraph(paragraph)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)

doc5.save(os.path.join(base_dir, "Task5_Reflection.docx"))
print("All 5 Task docx files created successfully!")
